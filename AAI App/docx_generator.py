"""
DOCX document generation for JG DAC Assessment Tool.

Generates rating forms and the DAC Report by modifying template files.
Templates are copied from the original sample files in templates_docx/.

Key conventions from the templates:
  - Rating forms: 4-col scoring tables, each KBI = 5 rows
      Row 0: KBI title (black bg, all 4 merged)
      Row 1: LEVEL 1 / LEVEL 2 / LEVEL 3 / LEVEL 4 (gray)
      Row 2: Level descriptors — yellow (FFFF00) on selected column
      Row 3: Key Progression (merged, unchanged)
      Row 4: Evidence (merged, filled with AI evidence)
  - Report:
      Table 3: Summary scores (3 rows × 5 comps, middle row shows dark fills ≤ rating)
      Tables 5–9: DSP / WTP / LFDV / MaD / ET narrative sections
      Table 10: Executive summary
"""

import io
import re
import zipfile
import copy
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATES_DIR = Path("templates_docx")

# Mapping from uppercase comp name in templates → comp code
COMP_NAME_MAP = {
    "DRIVE SUPERIOR PERFORMANCE": "DSP",
    "WIN THROUGH PEOPLE": "WTP",
    "LEARN FROM DIFFERENT VIEWS": "LFDV",
    "MAKE A DIFFERENCE": "MaD",
    "ESTABLISH TRUST": "ET",
}

# Level descriptor highlight: selected level = yellow, unselected = white
SELECTED_FILL = "FFFF00"
UNSELECTED_FILL = "FFFFFF"

# Report summary table: level bar colors
BAR_COLORS = {1: "EE0000", 2: "FFC000", 3: "92D050", 4: "00B050"}
DARK_FILL = "262A2D"

# Report comp narrative table indices (0-based from doc.tables)
REPORT_COMP_TABLE = {"DSP": 5, "WTP": 6, "LFDV": 7, "MaD": 8, "ET": 9}
# Summary table row groups (first row of each 3-row group, 0-based)
REPORT_SUMMARY_ROWS = {"DSP": 2, "WTP": 5, "LFDV": 8, "MaD": 11, "ET": 14}

ACTIVITY_TEMPLATE = {
    "ap":  "ap_rating_form.docx",
    "ge":  "ge_rating_form.docx",
    "cr":  "cr_rating_form.docx",
    "cbi": "cbi_guide.docx",
    "neg": "neg_rating_form.docx",  # First Loyalty — may not exist
}

# ────────────────────────────────────────────────────────────────────────────
# Low-level XML helpers
# ────────────────────────────────────────────────────────────────────────────

def get_unique_cells(row):
    """Return de-duplicated cells (merged cells appear once per row in python-docx)."""
    seen, out = set(), []
    for cell in row.cells:
        cid = id(cell._tc)
        if cid not in seen:
            seen.add(cid)
            out.append(cell)
    return out


def set_cell_fill(cell, fill_hex: str):
    """Set cell background shading.  fill_hex: 6-char hex, e.g. 'FFFF00'."""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)


def replace_cell_text(cell, new_text: str, prefix: str = ""):
    """
    Replace all text in a cell.
    If prefix is given the result is: "{prefix}\n{new_text}".
    Preserves the formatting of the first run.
    """
    full_text = f"{prefix}\n{new_text}" if prefix else new_text

    # Save formatting from first run of first paragraph
    first_rpr = None
    for para in cell.paragraphs:
        for run in para.runs:
            if run._r.find(qn("w:rPr")) is not None:
                first_rpr = copy.deepcopy(run._r.find(qn("w:rPr")))
            break
        if first_rpr is not None:
            break

    # Clear all paragraphs
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
        # remove extra elements (tabs, breaks, etc.)
        for child in list(para._p):
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag in ("r", "hyperlink", "bookmarkStart", "bookmarkEnd", "proofErr"):
                para._p.remove(child)

    # Write new text to first paragraph, first run
    para = cell.paragraphs[0]
    run = para.add_run(full_text)
    if first_rpr is not None:
        run._r.insert(0, copy.deepcopy(first_rpr))


def update_header_table(doc: Document, candidate_name: str, date: str, assessors: str):
    """Update Table 0: participant name, date, assessors."""
    t = doc.tables[0]
    # Row 0, col 0 → participant name
    cell_name = get_unique_cells(t.rows[0])[0]
    replace_cell_text(cell_name, candidate_name, prefix="Participant name:")
    # Row 0, col 1 → assessors
    cells_r0 = get_unique_cells(t.rows[0])
    if len(cells_r0) > 1:
        replace_cell_text(cells_r0[1], assessors, prefix="Assessors' names:")
    # Row 1, col 0 → date
    cells_r1 = get_unique_cells(t.rows[1])
    replace_cell_text(cells_r1[0], f"Date: {date}")


# ────────────────────────────────────────────────────────────────────────────
# Rating form generation (AP / GE / CR)
# ────────────────────────────────────────────────────────────────────────────

def update_rating_form_comp_table(table, comp_code: str, comp_data: dict):
    """
    Modify one competency scoring table in a rating form.
    comp_data: {"kbis": [...], "overall": 2, "rationale": "..."}
    """
    kbis = comp_data.get("kbis", [])
    overall = comp_data.get("overall", 2)
    n_rows = len(table.rows)

    # ── Row 0: overall rating ──────────────────────────────────────────────
    cells_r0 = get_unique_cells(table.rows[0])
    # The second unique cell spans cols 2-3 and shows "IJ Demonstration: X"
    if len(cells_r0) >= 2:
        replace_cell_text(cells_r0[1], f"IJ Demonstration: {overall}")

    # ── KBI blocks (5 rows each, starting at row 2) ────────────────────────
    for kbi_idx, kbi in enumerate(kbis):
        base = 2 + kbi_idx * 5
        if base + 4 >= n_rows:
            break

        rating = kbi.get("suggested_rating", overall)
        evidence = kbi.get("evidence", "")

        # Row base+2 = level descriptor row (4 unique cells = one per level)
        desc_row = table.rows[base + 2]
        desc_cells = get_unique_cells(desc_row)
        for col_idx, dcell in enumerate(desc_cells):
            level = col_idx + 1          # levels are 1-indexed
            fill = SELECTED_FILL if level == rating else UNSELECTED_FILL
            set_cell_fill(dcell, fill)

        # Row base+4 = evidence row (merged → 1 unique cell)
        ev_row = table.rows[base + 4]
        ev_cells = get_unique_cells(ev_row)
        if ev_cells:
            replace_cell_text(ev_cells[0], evidence,
                              prefix="Evidence / Additional Observations or Comments:")


def generate_rating_form(activity_code: str, candidate: dict, confirmed_data: dict) -> io.BytesIO:
    """
    Generate a rating form DOCX for the given activity.
    Returns a BytesIO of the document.
    """
    tpl_name = ACTIVITY_TEMPLATE.get(activity_code)
    if not tpl_name:
        raise ValueError(f"No template for activity: {activity_code}")

    tpl_path = TEMPLATES_DIR / tpl_name
    if not tpl_path.exists():
        raise FileNotFoundError(f"Template not found: {tpl_path}")

    doc = Document(str(tpl_path))

    # ── Header table ──────────────────────────────────────────────────────
    update_header_table(
        doc,
        candidate_name=candidate.get("name", ""),
        date=candidate.get("assessment_date", ""),
        assessors=candidate.get("assessors", ""),
    )

    # ── Scoring tables (index 4 onward) ──────────────────────────────────
    comps_data = confirmed_data.get("competencies", {})
    for t_idx in range(4, len(doc.tables)):
        table = doc.tables[t_idx]
        if not table.rows:
            continue
        # Identify comp from header cell (row 0, col 0)
        header_cells = get_unique_cells(table.rows[0])
        if not header_cells:
            continue
        header_text = header_cells[0].text.strip().upper()
        comp_code = COMP_NAME_MAP.get(header_text)
        if comp_code and comp_code in comps_data:
            update_rating_form_comp_table(table, comp_code, comps_data[comp_code])

    # ── Name replacement in raw XML (handles text boxes + footer) ────────
    buf = io.BytesIO()
    doc.save(buf)
    buf = _replace_name_in_docx(buf, "Motas, Elaine Dianne C.", candidate.get("name", ""))
    buf = _replace_name_in_docx(buf, "Elaine Motas", candidate.get("name", ""))
    buf = _replace_text_in_docx(buf, "June 10, 2026", candidate.get("assessment_date", ""))
    return buf


# ────────────────────────────────────────────────────────────────────────────
# CBI guide generation
# ────────────────────────────────────────────────────────────────────────────

def generate_cbi_guide(candidate: dict, confirmed_data: dict) -> io.BytesIO:
    """
    Generate the CBI interview guide DOCX.
    Each of the 5 competency tables gets AI evidence filled into the notes column.
    """
    tpl_path = TEMPLATES_DIR / "cbi_guide.docx"
    doc = Document(str(tpl_path))
    comps_data = confirmed_data.get("competencies", {})

    for table in doc.tables:
        if not table.rows:
            continue
        header_cells = get_unique_cells(table.rows[0])
        if not header_cells:
            continue
        # CBI header: "Leadership Brand: \nDRIVE SUPERIOR PERFORMANCE"
        header_text = header_cells[0].text.upper()
        comp_code = None
        for key, code in COMP_NAME_MAP.items():
            if key in header_text:
                comp_code = code
                break
        if not comp_code or comp_code not in comps_data:
            continue

        comp = comps_data[comp_code]
        kbis = comp.get("kbis", [])
        overall = comp.get("overall", 2)

        # Build combined notes text from all KBIs
        notes_lines = [f"Overall rating: {overall}", ""]
        for kbi in kbis:
            notes_lines.append(f"KBI {kbi['n']} (Rating: {kbi.get('suggested_rating', '?')})")
            notes_lines.append(kbi.get("evidence", ""))
            notes_lines.append("")
        notes_text = "\n".join(notes_lines).strip()

        # Fill col 2 (Interview Notes) of rows 1 and 2
        for ri in [1, 2]:
            if ri < len(table.rows):
                row_cells = get_unique_cells(table.rows[ri])
                if len(row_cells) >= 3:
                    # Only write to row 1; row 2 gets rationale
                    if ri == 1:
                        replace_cell_text(row_cells[2], notes_text)
                    else:
                        replace_cell_text(row_cells[2], comp.get("rationale", ""))

    buf = io.BytesIO()
    doc.save(buf)
    buf = _replace_name_in_docx(buf, "Motas, Elaine Dianne C.", candidate.get("name", ""))
    buf = _replace_name_in_docx(buf, "Elaine Motas", candidate.get("name", ""))
    buf = _replace_text_in_docx(buf, "June 10, 2026", candidate.get("assessment_date", ""))
    return buf


# ────────────────────────────────────────────────────────────────────────────
# DAC Report generation
# ────────────────────────────────────────────────────────────────────────────

def _update_report_summary_table(doc: Document, consolidated: dict):
    """Update Table 3 (summary of scores): ratings and bar fills."""
    t = doc.tables[3]

    for comp_code, start_row in REPORT_SUMMARY_ROWS.items():
        rating = consolidated.get(comp_code, 0)
        if not rating:
            continue

        for offset in range(3):
            ri = start_row + offset
            if ri >= len(t.rows):
                continue
            row_cells = get_unique_cells(t.rows[ri])
            if len(row_cells) < 2:
                continue

            # Col 1 = rating number (all 3 rows)
            replace_cell_text(row_cells[1], str(rating))

            # Middle row (offset==1): update level fills
            if offset == 1:
                level_colors = {
                    1: BAR_COLORS[1],   # EE0000
                    2: BAR_COLORS[2],   # FFC000
                    3: BAR_COLORS[3],   # 92D050
                    4: BAR_COLORS[4],   # 00B050
                }
                for level in range(1, 5):
                    col_idx = level + 1   # cols 2-5 correspond to levels 1-4
                    if col_idx < len(row_cells):
                        fill = DARK_FILL if level <= rating else level_colors[level]
                        set_cell_fill(row_cells[col_idx], fill)


def _update_report_comp_section(doc: Document, comp_code: str, comp_data: dict,
                                 activity_evidence: list[str], cbi_evidence: str):
    """
    Update one competency narrative table (Tables 5-9).
    activity_evidence: list of evidence strings from AP/GE/CR
    cbi_evidence: SBII text from CBI
    """
    t_idx = REPORT_COMP_TABLE.get(comp_code)
    if t_idx is None or t_idx >= len(doc.tables):
        return

    table = doc.tables[t_idx]
    overall = comp_data.get("overall", 2)

    # Row 0, col 2: overall rating
    r0_cells = get_unique_cells(table.rows[0])
    if len(r0_cells) >= 3:
        replace_cell_text(r0_cells[2], str(overall))

    # Row 4: behavioral observations from structured exercises
    if len(table.rows) > 4:
        r4_cells = get_unique_cells(table.rows[4])
        if r4_cells:
            kbi_bullets = []
            for kbi in comp_data.get("kbis", []):
                kbi_bullets.append(f"{kbi['title']} (Rating: {kbi.get('suggested_rating', '?')})\n{kbi.get('evidence', '')}")
            observations = "\n\n".join(activity_evidence) if activity_evidence else "\n\n".join(kbi_bullets)
            replace_cell_text(r4_cells[0],
                              f"The following behaviors were displayed:\n{observations}")

    # Row 6: SBII interview notes
    if len(table.rows) > 6 and cbi_evidence:
        r6_cells = get_unique_cells(table.rows[6])
        if r6_cells:
            replace_cell_text(r6_cells[0],
                              f"Follow the Situation-Behavior-Impact-Intent (SBII) format\n{cbi_evidence}")


def generate_dac_report(candidate: dict, all_confirmed: dict, consolidated: dict) -> io.BytesIO:
    """
    Generate the full DAC Report DOCX.
    all_confirmed: {activity_code: {competencies: {comp_code: comp_data}}}
    consolidated: {comp_code: max_rating}
    """
    tpl_path = TEMPLATES_DIR / "dac_report.docx"
    doc = Document(str(tpl_path))

    # ── Table 1: assessors ───────────────────────────────────────────────
    t1 = doc.tables[1]
    # Row 1, col 1 → Lead Assessor
    cells_r1 = get_unique_cells(t1.rows[1])
    if len(cells_r1) > 1:
        assessors_raw = candidate.get("assessors", "")
        parts = [a.strip() for a in assessors_raw.replace("/", ",").split(",") if a.strip()]
        lead = parts[0] if parts else assessors_raw
        replace_cell_text(cells_r1[1],
                          f"Lead Assessor:\n\n{lead}\nSHL Certified Assessor")
    # Row 2, col 1 → Co-Assessors
    cells_r2 = get_unique_cells(t1.rows[2])
    if len(cells_r2) > 1:
        co = ", ".join(parts[1:]) if len(parts) > 1 else ""
        replace_cell_text(cells_r2[1],
                          f"Co-Assessors:\n\n{co}\nSHL Certified Assessor" if co
                          else "Co-Assessors:\n\n—")

    # ── Table 3: summary scores ──────────────────────────────────────────
    _update_report_summary_table(doc, consolidated)

    # ── Tables 5-9: narrative per competency ─────────────────────────────
    COMP_ORDER = ["DSP", "WTP", "LFDV", "MaD", "ET"]
    for comp_code in COMP_ORDER:
        # Collect evidence from all activities that measured this comp
        activity_evidence = []
        cbi_evidence = ""
        for act_code, act_confirmed in all_confirmed.items():
            comp_data_act = act_confirmed.get("competencies", {}).get(comp_code)
            if not comp_data_act:
                continue
            if act_code == "cbi":
                cbi_lines = [f"KBI {k['n']}: {k.get('evidence', '')}"
                             for k in comp_data_act.get("kbis", [])]
                cbi_evidence = "\n\n".join(cbi_lines)
            else:
                lines = [f"KBI {k['n']} — {k['title']} (Rating: {k.get('suggested_rating','?')})\n{k.get('evidence','')}"
                         for k in comp_data_act.get("kbis", [])]
                activity_evidence.append("\n\n".join(lines))

        # Use consolidated data for overall rating display
        best_comp_data = None
        for act_code in ["cbi", "ap", "ge", "cr", "neg"]:
            d = all_confirmed.get(act_code, {}).get("competencies", {}).get(comp_code)
            if d:
                if best_comp_data is None or d.get("overall", 0) > best_comp_data.get("overall", 0):
                    best_comp_data = d

        if best_comp_data is None:
            best_comp_data = {"overall": consolidated.get(comp_code, 0), "kbis": []}

        _update_report_comp_section(doc, comp_code, best_comp_data,
                                    activity_evidence, cbi_evidence)

    # ── Table 10: executive summary ──────────────────────────────────────
    if len(doc.tables) > 10:
        t10 = doc.tables[10]
        # Row 1: career aspiration
        if len(t10.rows) > 1:
            cells = get_unique_cells(t10.rows[1])
            if cells:
                # Pull from CBI data if available
                cbi = all_confirmed.get("cbi", {})
                aspiration = cbi.get("career_aspiration", "")
                if not aspiration:
                    aspiration = "[Please add career aspiration from CBI notes]"
                replace_cell_text(cells[0], aspiration)
        # Row 3: overall summary & recommendation
        if len(t10.rows) > 3:
            cells = get_unique_cells(t10.rows[3])
            if cells:
                summary = _build_executive_summary(candidate, consolidated)
                replace_cell_text(cells[0], summary)

    # ── Save and name-swap ───────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    cand_name = candidate.get("name", "")
    buf = _replace_name_in_docx(buf, "Elaine Motas", cand_name)
    buf = _replace_name_in_docx(buf, "Motas, Elaine Dianne C.", cand_name)
    buf = _replace_text_in_docx(buf, "June 10, 2026", candidate.get("assessment_date", ""))
    buf = _replace_text_in_docx(buf, "Sharlyn Sanclaria / Ana Roberta", candidate.get("assessors", ""))
    return buf


def _build_executive_summary(candidate: dict, consolidated: dict) -> str:
    """Build a placeholder executive summary text."""
    name = candidate.get("name", "The participant")
    lines = [
        f"Integrating the results of the Assessment Center activities, "
        f"{name} demonstrated the following overall performance:\n"
    ]
    LEVEL_LABELS = {
        1: "Does Not Demonstrate",
        2: "Inconsistently Demonstrates",
        3: "Consistently Demonstrates",
        4: "Exceeds Expectations",
    }
    COMP_NAMES = {
        "DSP": "Drive Superior Performance",
        "WTP": "Win Through People",
        "LFDV": "Learn from Different Views",
        "MaD": "Make a Difference",
        "ET": "Establish Trust",
    }
    for comp_code, score in consolidated.items():
        label = LEVEL_LABELS.get(score, "")
        comp_name = COMP_NAMES.get(comp_code, comp_code)
        lines.append(f"• {comp_name}: {score} — {label}")
    lines.append("\n[Please add detailed narrative and recommendation here]")
    return "\n".join(lines)


# ────────────────────────────────────────────────────────────────────────────
# Raw XML text replacement (for name in text boxes / footers)
# ────────────────────────────────────────────────────────────────────────────

def _replace_text_in_docx(buf: io.BytesIO, old: str, new: str) -> io.BytesIO:
    """Replace a plain string in all XML parts of a DOCX buffer."""
    if not old or old == new:
        return buf
    buf.seek(0)
    out = io.BytesIO()
    with zipfile.ZipFile(buf, "r") as zin, zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.endswith(".xml") or item.filename.endswith(".rels"):
                text = data.decode("utf-8", errors="replace")
                if old in text:
                    text = text.replace(old, new)
                    data = text.encode("utf-8")
            zout.writestr(item, data)
    out.seek(0)
    return out


def _replace_name_in_docx(buf: io.BytesIO, old_name: str, new_name: str) -> io.BytesIO:
    """
    Replace a name that may be split across multiple <w:t> XML elements.
    Uses both direct replace and a regex that handles split runs.
    """
    if not old_name or old_name == new_name:
        return buf

    # First try direct replace (handles single-run names)
    buf = _replace_text_in_docx(buf, old_name, new_name)

    # Handle split across runs: e.g. "Elaine</w:t>...</w:t>Motas"
    # Build a regex pattern that matches the name across run boundaries
    words = old_name.split()
    if len(words) > 1:
        buf.seek(0)
        out = io.BytesIO()
        with zipfile.ZipFile(buf, "r") as zin, zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename.endswith(".xml"):
                    text = data.decode("utf-8", errors="replace")
                    # Build flexible pattern
                    pattern = r"(<w:t[^>]*>)" + r"(</w:t>.*?<w:t[^>]*>)*".join(
                        re.escape(w) for w in words
                    ) + r"(</w:t>)"
                    if re.search(pattern, text, re.DOTALL):
                        text = re.sub(
                            pattern,
                            lambda m: m.group(0).split(words[0])[0] + new_name + m.group(0).rsplit(words[-1])[-1],
                            text, flags=re.DOTALL
                        )
                        data = text.encode("utf-8")
                zout.writestr(item, data)
        out.seek(0)
        return out

    return buf


# ────────────────────────────────────────────────────────────────────────────
# Public dispatch function
# ────────────────────────────────────────────────────────────────────────────

def generate_activity_document(activity_code: str, candidate: dict, confirmed_data: dict) -> io.BytesIO:
    """
    Generate the appropriate DOCX for a single activity.
    Returns BytesIO ready for streaming.
    """
    if activity_code == "cbi":
        return generate_cbi_guide(candidate, confirmed_data)
    else:
        return generate_rating_form(activity_code, candidate, confirmed_data)
